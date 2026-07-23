import io

import docx
import fitz  # PyMuPDF
from app.core.exceptions import EmptyResumeTextError, UnsupportedFileTypeError

SUPPORTED_CONTENT_TYPES = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
}


def detect_file_kind(filename: str, content_type: str | None) -> str:
    """
    Prefer content_type from the multipart upload, fall back to file extension.
    Never trust the client blindly — validate against an explicit allowlist.
    """
    if content_type in SUPPORTED_CONTENT_TYPES:
        return SUPPORTED_CONTENT_TYPES[content_type]

    lowered = filename.lower()
    if lowered.endswith(".pdf"):
        return "pdf"
    if lowered.endswith(".docx"):
        return "docx"

    raise UnsupportedFileTypeError(
        f"Unsupported file type for '{filename}'. Only PDF and DOCX are supported.",
        details={"filename": filename, "content_type": content_type},
    )


def extract_text_from_pdf(file_bytes: bytes) -> str:
    text_parts: list[str] = []
    with fitz.open(stream=file_bytes, filetype="pdf") as document:
        for page in document:
            text_parts.append(page.get_text())
    return "\n".join(text_parts).strip()


def extract_text_from_docx(file_bytes: bytes) -> str:
    document = docx.Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]

    # Tables often hold skills/experience in resume templates — don't skip them.
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)

    return "\n".join(paragraphs).strip()


def extract_resume_text(
    filename: str, content_type: str | None, file_bytes: bytes
) -> str:
    kind = detect_file_kind(filename, content_type)

    text = (
        extract_text_from_pdf(file_bytes)
        if kind == "pdf"
        else extract_text_from_docx(file_bytes)
    )

    if not text or len(text) < 20:
        raise EmptyResumeTextError(
            "Could not extract readable text from the uploaded resume. "
            "It may be a scanned image or corrupted file.",
            details={"filename": filename},
        )

    return text
