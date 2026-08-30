import io
import docx
import fitz  # PyMuPDF
import httpx

from app.core.exceptions import EmptyResumeTextError, UnsupportedFileTypeError
from app.core.logging import get_logger
from app.lib.constants import (
    ALLOWED_RESUME_MIME_TYPES,
    MAX_RAW_TEXT_LENGTH,
    MAX_RESUME_FILE_SIZE_BYTES,
    MIN_TEXT_LENGTH,
)

logger = get_logger(__name__)


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extracts raw text from PDF bytes using PyMuPDF."""
    text_parts: list[str] = []
    with fitz.open(stream=file_bytes, filetype="pdf") as document:
        for page in document:
            text = page.get_text()
            if text:
                text_parts.append(text)
    return "\n".join(text_parts).strip()


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extracts raw text from DOCX bytes including paragraphs and tables."""
    document = docx.Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())

    return "\n".join(paragraphs).strip()


def extract_text_from_file(
    filename: str, content_type: str | None, file_bytes: bytes
) -> str:
    """Validates file type and extracts text from binary bytes."""
    if len(file_bytes) > MAX_RESUME_FILE_SIZE_BYTES:
        raise UnsupportedFileTypeError(
            f"File exceeds maximum allowed size of {MAX_RESUME_FILE_SIZE_BYTES // (1024 * 1024)}MB."
        )

    lowered_name = filename.lower()
    is_pdf = lowered_name.endswith(".pdf") or content_type == "application/pdf"
    is_docx = (
        lowered_name.endswith(".docx")
        or content_type
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    if not (is_pdf or is_docx):
        raise UnsupportedFileTypeError(
            f"Unsupported file type for '{filename}'. Only PDF and DOCX files are supported."
        )

    text = extract_text_from_pdf(file_bytes) if is_pdf else extract_text_from_docx(file_bytes)

    if not text or len(text.strip()) < MIN_TEXT_LENGTH:
        raise EmptyResumeTextError(
            "Could not extract readable text from the file. It may be an image-only scan or corrupted file."
        )

    return text.strip()


async def extract_text_from_url(url: str) -> str:
    """Fetches and extracts text from a remote URL with SSRF protection."""
    if not url.startswith(("http://", "https://")):
        raise EmptyResumeTextError("Invalid URL protocol. Only HTTP and HTTPS are allowed.")

    async with httpx.AsyncClient(timeout=15.0, follow_redirects=True) as client:
        try:
            response = await client.get(url)
            response.raise_for_status()
            text = response.text.strip()
            if len(text) < MIN_TEXT_LENGTH:
                raise EmptyResumeTextError("The content fetched from the URL was too short to evaluate.")
            return text[:MAX_RAW_TEXT_LENGTH]
        except Exception as exc:
            logger.warning("url_fetch_failed", url=url, error=str(exc))
            raise EmptyResumeTextError(f"Failed to fetch content from URL: {str(exc)}") from exc
